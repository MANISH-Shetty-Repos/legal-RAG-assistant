"""
Document Loaders — PDF, DOCX, TXT, Markdown
Extracts text content and page-level metadata from various document formats.
"""

from pathlib import Path
from datetime import datetime, timezone
from dataclasses import dataclass, field

from loguru import logger


@dataclass
class PageContent:
    """Represents a single page/section of extracted text."""

    text: str
    page_number: int
    metadata: dict = field(default_factory=dict)


@dataclass
class LoadedDocument:
    """Represents a fully loaded document with all pages."""

    filename: str
    file_path: str
    file_type: str
    upload_date: str
    pages: list[PageContent] = field(default_factory=list)
    total_pages: int = 0

    def get_full_text(self) -> str:
        """Concatenate all page texts."""
        return "\n\n".join(page.text for page in self.pages)


class PDFLoader:
    """Load and extract text from PDF files using pdfplumber (primary) or pypdf (fallback)."""

    @staticmethod
    def load(file_path: str) -> LoadedDocument:
        """Extract text and metadata from a PDF file."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"PDF file not found: {file_path}")

        pages: list[PageContent] = []

        try:
            # Primary: pdfplumber (better table/layout extraction)
            import pdfplumber

            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        pages.append(
                            PageContent(
                                text=text.strip(),
                                page_number=i + 1,
                                metadata={"source": "pdfplumber"},
                            )
                        )
                logger.info(
                    f"Loaded PDF with pdfplumber: {path.name} ({len(pages)} pages)"
                )

        except ImportError:
            # Fallback: pypdf
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(
                        PageContent(
                            text=text.strip(),
                            page_number=i + 1,
                            metadata={"source": "pypdf"},
                        )
                    )
            logger.info(f"Loaded PDF with pypdf: {path.name} ({len(pages)} pages)")

        return LoadedDocument(
            filename=path.name,
            file_path=str(path.absolute()),
            file_type="pdf",
            upload_date=datetime.now(timezone.utc).isoformat(),
            pages=pages,
            total_pages=len(pages),
        )


class DOCXLoader:
    """Load and extract text from DOCX files."""

    @staticmethod
    def load(file_path: str) -> LoadedDocument:
        """Extract text from a DOCX file."""
        from docx import Document as DocxDocument

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        doc = DocxDocument(file_path)
        full_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())

        # DOCX doesn't have native pages, treat entire doc as page 1
        pages = []
        if full_text.strip():
            pages.append(
                PageContent(
                    text=full_text.strip(),
                    page_number=1,
                    metadata={
                        "source": "python-docx",
                        "paragraph_count": len(doc.paragraphs),
                    },
                )
            )

        logger.info(f"Loaded DOCX: {path.name} ({len(doc.paragraphs)} paragraphs)")

        return LoadedDocument(
            filename=path.name,
            file_path=str(path.absolute()),
            file_type="docx",
            upload_date=datetime.now(timezone.utc).isoformat(),
            pages=pages,
            total_pages=1,
        )


class TextLoader:
    """Load plain text and markdown files."""

    @staticmethod
    def load(file_path: str) -> LoadedDocument:
        """Extract text from a TXT or Markdown file."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Text file not found: {file_path}")

        text = path.read_text(encoding="utf-8")
        file_type = "markdown" if path.suffix.lower() in (".md", ".markdown") else "txt"

        pages = []
        if text.strip():
            pages.append(
                PageContent(
                    text=text.strip(),
                    page_number=1,
                    metadata={"source": "text_loader", "char_count": len(text)},
                )
            )

        logger.info(f"Loaded {file_type.upper()}: {path.name} ({len(text)} chars)")

        return LoadedDocument(
            filename=path.name,
            file_path=str(path.absolute()),
            file_type=file_type,
            upload_date=datetime.now(timezone.utc).isoformat(),
            pages=pages,
            total_pages=1,
        )


# --- Unified Loader Interface ---

# Map of file extensions to their loader classes
LOADER_MAP = {
    ".pdf": PDFLoader,
    ".docx": DOCXLoader,
    ".txt": TextLoader,
    ".md": TextLoader,
    ".markdown": TextLoader,
}

SUPPORTED_EXTENSIONS = set(LOADER_MAP.keys())


def load_document(file_path: str) -> LoadedDocument:
    """
    Load a document using the appropriate loader based on file extension.

    Args:
        file_path: Path to the document file

    Returns:
        LoadedDocument with extracted text and metadata

    Raises:
        ValueError: If the file extension is not supported
        FileNotFoundError: If the file does not exist
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext not in LOADER_MAP:
        raise ValueError(
            f"Unsupported file type: '{ext}'. "
            f"Supported formats: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    loader = LOADER_MAP[ext]
    return loader.load(file_path)


def load_directory(dir_path: str) -> list[LoadedDocument]:
    """
    Load all supported documents from a directory.

    Args:
        dir_path: Path to the directory

    Returns:
        List of LoadedDocument objects
    """
    path = Path(dir_path)
    if not path.is_dir():
        raise NotADirectoryError(f"Not a directory: {dir_path}")

    documents = []
    for file_path in sorted(path.iterdir()):
        if file_path.suffix.lower() in SUPPORTED_EXTENSIONS:
            try:
                doc = load_document(str(file_path))
                documents.append(doc)
            except Exception as e:
                logger.error(f"Failed to load {file_path.name}: {e}")

    logger.info(f"Loaded {len(documents)} documents from {dir_path}")
    return documents
